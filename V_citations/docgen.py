import pandas as pd
from pandas import DataFrame, Series
import numpy as np
from datetime import date
import uuid
from docx import Document
import re


replacements = {
    "{collection}": "collection",
    "{name}": "name",
    "{art_date}": "productOptionDescription2",
    "{medium}": "additionalInfoDescription3",
    "{dimension}": "additionalInfoDescription4",
    "{exhibit_place}": "additionalInfoDescription5",
    "{town}": "productOptionDescription3",
    "{photo_credit}": "additionalInfoDescription2"
}

to_remove = ['<p>', '</p>', '<strong>', '</strong>', '<span>', '</span>', '<tr>', '</tr>', '<td>', '</td>', '<li>', '<ul>', '</li>', '</ul>'
             '&nbsp', '\n', '\\n', '/', '<br>', '<br/>', '</br>', ';']


def std_value(val, val_type):
    new_val = str(val)

    for tr in to_remove:
        new_val = new_val.replace(tr, '')

    if val_type == "{collection}":
        new_val = new_val.title()

    if val_type == "{name}":
        new_val = new_val.title()

    if val_type == "{dimension}":
        new_val = new_val.lower()

    return new_val.strip()


# paragraph format : {collection}, {name}, {art_date}, {medium}, {dimension}, {exhibit_place}, {town}, © photo : {photo_credit}


def prepare_run(p, placeholder_key, possible_value):

    run = p.add_run()
    val = std_value(possible_value, placeholder_key)

    if val == "":
        return

    if placeholder_key == "{collection}":
        run.text = val
    elif placeholder_key == "{name}":
        run.text = ", " + val
        run.italic = True
    else:
        run.text = ", " + val


def doc_gen(datafile) -> str:
    res = Document()

    # test if file exist
    try:
        open(datafile)
    except FileNotFoundError:
        raise FileNotFoundError("The file could not be opened")

    print("Will read csv file !")
    df: DataFrame = pd.read_csv(datafile, sep=',', keep_default_na=False)
    print("Has read csv file !")
    
    print("file read : ", df.columns)
    print("file read : ", df.head())
    print("file nb lines : ", len(df))

    ## check that all colmuns are present in file 
    errors = []
    for key in replacements.values():
        if key not in df.columns:
            errors.append(key)
    
    if len(errors) > 0:
        raise KeyError(f"Keys {errors} not found in csv file")

    for index, row in df.iterrows():
        p = res.add_paragraph("")

        for placeholder_key, data_col in replacements.items():
            possible_value = row[data_col]
            prepare_run(p, placeholder_key, possible_value)

    res.save("./generated_citation.docx")
    return "./generated_citation.docx"
